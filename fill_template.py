# -*- coding: utf-8 -*-
"""
fill_template.py
================
Preenche templates .docx de plano alimentar veterinario (Apsirtos) e retorna
os bytes do documento pronto para conversao via LibreOffice.
"""

from __future__ import annotations

import io
import os
from typing import List, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Caminho para a assinatura digital (relativo a este arquivo)
_ASSETS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "assets")
_SIG_PATH   = os.path.join(_ASSETS_DIR, "assinatura.png")


class _C:
    TEAL_DARK   = RGBColor(0x1A, 0x52, 0x76)
    TEAL_MED    = RGBColor(0x2E, 0x86, 0xC1)
    TEAL_LIGHT  = RGBColor(0xEB, 0xF5, 0xFB)
    TABLE_HDR   = "1A5276"
    WHITE       = "FFFFFF"
    TEXT_DARK   = RGBColor(0x1C, 0x28, 0x33)


_TBLPR_ORDER = [
    "tblStyle", "tblpPr", "tblOverlap", "bidiVisual",
    "tblStyleRowBandSize", "tblStyleColBandSize",
    "tblW", "jc", "tblCellSpacing", "tblInd",
    "tblBorders", "shd", "tblLayout", "tblCellMar",
    "tblLook", "tblCaption", "tblPrChange",
]

_TCPR_ORDER = [
    "cnfStyle", "tcW", "gridSpan", "vMerge",
    "tcBorders", "shd", "noWrap", "tcMar",
    "textDirection", "tcFitText", "vAlign",
    "hideMark", "headers", "cellIns", "cellDel", "trPr",
]


def _insert_ordered(parent, new_el, order: list) -> None:
    ns_tag = new_el.tag
    local  = ns_tag.split("}")[1] if "}" in ns_tag else ns_tag
    for old in parent.findall(ns_tag):
        parent.remove(old)
    new_pos = order.index(local) if local in order else len(order)
    insert_before = None
    for child in list(parent):
        child_local = child.tag.split("}")[1] if "}" in child.tag else child.tag
        child_pos = order.index(child_local) if child_local in order else len(order)
        if child_pos > new_pos:
            insert_before = child
            break
    if insert_before is not None:
        insert_before.addprevious(new_el)
    else:
        parent.append(new_el)


def _set_cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    _insert_ordered(tcPr, shd, _TCPR_ORDER)


def _set_cell_width(cell, width_dxa: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    _insert_ordered(tcPr, tcW, _TCPR_ORDER)


def _set_cell_margins(cell, top=80, bottom=80, left=140, right=140) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left),
                      ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    _insert_ordered(tcPr, tcMar, _TCPR_ORDER)


def _set_cell_borders(cell, color: str = "CCCCCC") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    borders_el = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        borders_el.append(b)
    _insert_ordered(tcPr, borders_el, _TCPR_ORDER)


def _cell_valign(cell, val: str = "center") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), val)
    _insert_ordered(tcPr, vAlign, _TCPR_ORDER)


def _set_row_height(row, twips: int, rule: str = "atLeast") -> None:
    trPr = row._tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(twips))
    trH.set(qn("w:hRule"), rule)
    trPr.append(trH)


def _para_space(para, before=0, after=0) -> None:
    pPr = para._element.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"),  str(after))


def _set_table_width(table, width_dxa: int) -> None:
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(width_dxa))
    tblW.set(qn("w:type"), "dxa")
    _insert_ordered(tblPr, tblW, _TBLPR_ORDER)


def _remove_table_borders(table) -> None:
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        return
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "none")
        b.set(qn("w:sz"),    "0")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "auto")
        borders.append(b)
    _insert_ordered(tblPr, borders, _TBLPR_ORDER)


def _replace_in_paragraph(para, placeholder: str, value: str) -> None:
    for run in para.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, value)
            return
    full = "".join(r.text for r in para.runs)
    if placeholder in full:
        new_text = full.replace(placeholder, value)
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""


def _replace_all_placeholders(doc: Document, replacements: dict) -> None:
    def process_para(p):
        for key, val in replacements.items():
            ph = "{{" + key + "}}"
            if ph in p.text:
                _replace_in_paragraph(p, ph, str(val))

    for p in doc.paragraphs:
        process_para(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_para(p)

    # Escaneia rodape de secao (footer fixo com {{data}})
    for section in doc.sections:
        try:
            footer = section.footer
            for p in footer.paragraphs:
                process_para(p)
            for table in footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            process_para(p)
        except Exception:
            pass


def _build_ingredients_table(
    doc: Document,
    ingredients: List[Tuple[str, str]],
) -> object:
    CONTENT_W = 9638
    COL_INGR  = int(CONTENT_W * 0.745)
    COL_QTD   = CONTENT_W - COL_INGR

    table = doc.add_table(rows=0, cols=2)
    _set_table_width(table, CONTENT_W)
    _remove_table_borders(table)

    # Cabecalho
    hdr_row = table.add_row()
    _set_row_height(hdr_row, 640)
    for idx, (cell, header_text, col_w) in enumerate(zip(
        hdr_row.cells,
        ["Ingredientes", "Qtd"],
        [COL_INGR, COL_QTD],
    )):
        cell.text = ""
        _set_cell_bg(cell, _C.TABLE_HDR)
        _set_cell_borders(cell, _C.TABLE_HDR)
        _set_cell_width(cell, col_w)
        _set_cell_margins(cell, top=0, bottom=0, left=160, right=160)
        _cell_valign(cell, "center")
        p = cell.paragraphs[0]
        p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if idx == 1
                       else WD_ALIGN_PARAGRAPH.LEFT)
        _para_space(p, before=0, after=0)
        run = p.add_run(header_text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Linhas de ingredientes
    for i, (ingredient, qty) in enumerate(ingredients):
        bg = _C.WHITE
        row = table.add_row()
        _set_row_height(row, 700)

        c0 = row.cells[0]
        c0.text = ""
        _set_cell_bg(c0, bg)
        _set_cell_borders(c0, "D5E8F8")
        _set_cell_width(c0, COL_INGR)
        _set_cell_margins(c0, top=0, bottom=0, left=160, right=160)
        _cell_valign(c0, "center")
        p0 = c0.paragraphs[0]
        _para_space(p0, before=0, after=0)
        r0 = p0.add_run(str(ingredient))
        r0.font.name = "Arial"
        r0.font.size = Pt(11)
        r0.font.color.rgb = _C.TEXT_DARK

        c1 = row.cells[1]
        c1.text = ""
        _set_cell_bg(c1, bg)
        _set_cell_borders(c1, "D5E8F8")
        _set_cell_width(c1, COL_QTD)
        _set_cell_margins(c1, top=0, bottom=0, left=160, right=160)
        _cell_valign(c1, "center")
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _para_space(p1, before=0, after=0)
        r1 = p1.add_run(str(qty))
        r1.font.name = "Arial"
        r1.font.size = Pt(11)
        r1.font.color.rgb = _C.TEXT_DARK

    return table._tbl


def _add_signature_to_footer(doc: Document) -> None:
    """Insere a assinatura no rodapé, ACIMA da linha separadora.

    Estratégia: adiciona parágrafo com imagem na célula esquerda da tabela do
    rodapé, como primeiro parágrafo (antes do separador). O parágrafo é criado
    via left_cell.add_paragraph() (contexto correto do footer part para registrar
    o relacionamento da imagem), depois movido para antes do Para[0].
    A célula direita recebe vAlign=bottom para manter a data alinhada à base.

    Medidas: assinatura recortada = 710×200 px @ Inches(1.5) → ~1.07 cm de altura.
    Rodapé expandido para 2.7 cm (bottom_margin=3.5, footer_distance=0.8).
    """
    if not os.path.exists(_SIG_PATH):
        return

    for section in doc.sections:
        try:
            footer = section.footer
            if not footer.tables:
                continue

            left_cell  = footer.tables[0].rows[0].cells[0]
            right_cell = footer.tables[0].rows[0].cells[1]

            # 1. Adicionar parágrafo com imagem ao final da célula esquerda
            #    (contexto correto do footer part → add_picture registra rId certo)
            sig_para = left_cell.add_paragraph()
            _para_space(sig_para, before=0, after=60)
            sig_para.add_run().add_picture(_SIG_PATH, width=Inches(1.5))

            # 2. Mover para ANTES do Para[0] (linha separadora)
            first_para_el = left_cell.paragraphs[0]._element
            sig_para._element.getparent().remove(sig_para._element)
            first_para_el.addprevious(sig_para._element)

            # 3. Célula direita: vAlign=bottom para manter data alinhada à base
            tcPr = right_cell._tc.get_or_add_tcPr()
            existing_va = tcPr.find(qn("w:vAlign"))
            if existing_va is not None:
                tcPr.remove(existing_va)
            vAlign_el = OxmlElement("w:vAlign")
            vAlign_el.set(qn("w:val"), "bottom")
            tcPr.append(vAlign_el)

        except Exception:
            pass


def _add_signature_at_end(doc: Document) -> None:
    """Adiciona assinatura no rodapé (para exames e prescrição)."""
    _add_signature_to_footer(doc)


def _add_signature_before_pagebreak(doc: Document) -> None:
    """Adiciona assinatura no rodapé (para dieta)."""
    _add_signature_to_footer(doc)


def fill_template(
    docx_path: str,
    data: dict,
    ingredients: List[Tuple[str, str]],
    opcao_num: str,
) -> bytes:
    doc = Document(docx_path)

    replacements = dict(data)
    replacements["opcao_num"] = str(opcao_num)

    _replace_all_placeholders(doc, replacements)

    sample_table = None
    for t in doc.tables:
        try:
            if t.rows[0].cells[0].text.strip() == "Ingredientes":
                sample_table = t
                break
        except (IndexError, AttributeError):
            continue

    if sample_table is None:
        raise ValueError(
            "Tabela de ingredientes nao encontrada no template. "
            "Certifique-se de usar um template gerado por create_templates.py."
        )

    tbl_el = _build_ingredients_table(doc, ingredients)

    sample_tbl_el = sample_table._tbl
    sample_tbl_el.addprevious(tbl_el)
    sample_tbl_el.getparent().remove(sample_tbl_el)

    _add_signature_before_pagebreak(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def fill_exames(
    docx_path: str,
    data: dict,
    exames: List[str],
) -> bytes:
    """Preenche o template de solicitação de exames.

    Parameters
    ----------
    docx_path : str
        Caminho para template_exames.docx
    data : dict
        Mesmos campos do fill_template (pet_nome, tutor_nome, data, etc.)
    exames : List[str]
        Lista de exames solicitados — cada item vira um bullet point.
    """
    doc = Document(docx_path)

    _replace_all_placeholders(doc, data)

    # Localiza o parágrafo marcador {{exames}}
    marker = None
    for p in doc.paragraphs:
        if "{{exames}}" in p.text:
            marker = p
            break

    if marker is None:
        raise ValueError(
            "Marcador {{exames}} nao encontrado no template. "
            "Use um template gerado por create_templates.py."
        )

    # Insere os bullets antes do marcador (ordem inversa para manter sequência)
    TEXT_DARK = RGBColor(0x1C, 0x28, 0x33)
    for exame in exames:
        new_p = OxmlElement("w:p")
        # pPr — indentação bullet
        pPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "360")
        ind.set(qn("w:hanging"), "180")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "60")
        pPr.append(spacing)
        pPr.append(ind)
        new_p.append(pPr)
        # run com bullet + texto
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Arial")
        rFonts.set(qn("w:hAnsi"), "Arial")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "22")        # 11pt = 22 half-points
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), "1C2833")
        rPr.append(rFonts); rPr.append(sz); rPr.append(color_el)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = "• " + str(exame)
        r.append(t)
        new_p.append(r)
        marker._element.addprevious(new_p)

    # Remove o paragrafo marcador
    marker._element.getparent().remove(marker._element)

    _add_signature_at_end(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ─── Helpers internos para prescrição ────────────────────────────────────────

_PRESC_W        = 9638
_PRESC_HALF     = _PRESC_W // 2
_PRESC_TEAL_D   = "1A5276"
_PRESC_TEAL_M   = "2E86C1"
_PRESC_TEAL_L   = "EBF5FB"
_PRESC_WHITE_H  = "FFFFFF"
_PRESC_WHITE_R  = RGBColor(0xFF, 0xFF, 0xFF)
_PRESC_AED6F1   = RGBColor(0xAE, 0xD6, 0xF1)
_PRESC_DARK_R   = RGBColor(0x1C, 0x28, 0x33)
_PRESC_TEAL_DR  = RGBColor(0x1A, 0x52, 0x76)
_PRESC_TEAL_MR  = RGBColor(0x2E, 0x86, 0xC1)
_PRESC_GRAY_R   = RGBColor(0x55, 0x55, 0x55)


def _pw(table, dxa):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(dxa)); tblW.set(qn("w:type"), "dxa")
    _insert_ordered(tblPr, tblW, _TBLPR_ORDER)


def _prb(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None: return
    borders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),"none"); b.set(qn("w:sz"),"0")
        b.set(qn("w:space"),"0"); b.set(qn("w:color"),"auto")
        borders.append(b)
    _insert_ordered(tblPr, borders, _TBLPR_ORDER)


def _pcb(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
    shd.set(qn("w:fill"), hex_color)
    _insert_ordered(tcPr, shd, _TCPR_ORDER)


def _pcm(cell, top=80, bottom=80, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top",top),("left",left),("bottom",bottom),("right",right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"),"dxa")
        tcMar.append(el)
    _insert_ordered(tcPr, tcMar, _TCPR_ORDER)


def _pcbn(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    bel = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),"none"); b.set(qn("w:sz"),"0")
        b.set(qn("w:space"),"0"); b.set(qn("w:color"),"auto")
        bel.append(b)
    _insert_ordered(tcPr, bel, _TCPR_ORDER)


def _pcborder(cell, color, sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    bel = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),"single"); b.set(qn("w:sz"),sz)
        b.set(qn("w:space"),"0"); b.set(qn("w:color"),color)
        bel.append(b)
    _insert_ordered(tcPr, bel, _TCPR_ORDER)


def _pcw(cell, dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(dxa)); tcW.set(qn("w:type"),"dxa")
    _insert_ordered(tcPr, tcW, _TCPR_ORDER)


def _pr(para, text, bold=False, italic=False, size_pt=10, color=None):
    run = para.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.name = "Arial"; run.font.size = Pt(size_pt)
    if color: run.font.color.rgb = color


def _presc_header(doc):
    table = doc.add_table(rows=1, cols=1)
    _pw(table, _PRESC_W); _prb(table)
    cell = table.rows[0].cells[0]
    cell.text = ""
    _pcb(cell, _PRESC_TEAL_D)
    _pcm(cell, top=200, bottom=200, left=240, right=240)
    _pcbn(cell)
    p1 = cell.paragraphs[0]
    _pr(p1, 'Apsirtos Clinica Veterinaria', bold=True, size_pt=16, color=_PRESC_WHITE_R)
    _para_space(p1, before=0, after=50)
    p2 = cell.add_paragraph()
    _pr(p2, 'Isabelle Rizzo Assumpção  —  CRMV 48652/SP', size_pt=10, color=_PRESC_AED6F1)
    _para_space(p2, before=0, after=70)
    p3 = cell.add_paragraph()
    _pr(p3, 'PRESCRIÇÃO', bold=True, size_pt=11, color=_PRESC_WHITE_R)
    _para_space(p3, before=0, after=0)


def _presc_line(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6")
    bot.set(qn("w:space"),"1"); bot.set(qn("w:color"), _PRESC_TEAL_M)
    pBdr.append(bot); pPr.append(pBdr)
    _para_space(p, before=80, after=80)


def _presc_data(doc):
    table = doc.add_table(rows=1, cols=2)
    _pw(table, _PRESC_W); _prb(table)
    sections = [
        ("DADOS DO ANIMAL", [
            ("Nome",       "{{pet_nome}}"),
            ("Especie",    "{{pet_especie}}"),
            ("Raca",       "{{pet_raca}}"),
            ("Sexo",       "{{pet_sexo}}"),
            ("Nascimento", "{{pet_nascimento}}"),
        ]),
        ("DADOS DO RESPONSAVEL", [
            ("Nome",      "{{tutor_nome}}"),
            ("CPF",       "{{tutor_cpf}}"),
            ("Endereco",  "{{tutor_endereco}}"),
        ]),
    ]
    labels = {
        "DADOS DO ANIMAL": "DADOS DO ANIMAL",
        "DADOS DO RESPONSAVEL": "DADOS DO RESPONSÁVEL",
        "Especie": "Espécie", "Raca": "Raça", "Endereco": "Endereço",
    }
    for col_idx, (titulo_key, campos) in enumerate(sections):
        cell = table.rows[0].cells[col_idx]
        cell.text = ""
        _pcb(cell, _PRESC_TEAL_L)
        _pcborder(cell, _PRESC_TEAL_M)
        _pcw(cell, _PRESC_HALF)
        _pcm(cell, top=150, bottom=150, left=170, right=170)
        titulo_display = labels.get(titulo_key, titulo_key)
        p_title = cell.paragraphs[0]
        _pr(p_title, titulo_display, bold=True, size_pt=10, color=_PRESC_TEAL_DR)
        _para_space(p_title, before=0, after=90)
        for label_key, placeholder in campos:
            p = cell.add_paragraph()
            label_display = labels.get(label_key, label_key)
            _pr(p, f"{label_display}: ", bold=True, size_pt=9.5, color=_PRESC_DARK_R)
            _pr(p, placeholder, bold=False, size_pt=9.5, color=_PRESC_DARK_R)
            _para_space(p, before=0, after=20)


def fill_prescricao(
    docx_path: str,
    data: dict,
    pages: List[List[dict]],
) -> bytes:
    """Preenche o template de prescricao.

    Parameters
    ----------
    docx_path : str
        Caminho para template_prescricao.docx
    data : dict
        Campos do pet e tutor (pet_nome, pet_especie, pet_raca, pet_sexo,
        pet_nascimento, tutor_nome, tutor_cpf, tutor_endereco, data)
    pages : List[List[dict]]
        Lista de paginas. Cada pagina e uma lista de vias.
        Cada via: {"via": "USO ORAL", "medicamentos": [{"item": "...", "instrucao": "..."}]}
        A numeracao dos itens e sequencial atraves de todas as paginas e vias.

    Example
    -------
    pages = [
        [{"via": "USO ORAL", "medicamentos": [
            {"item": "Seren Snacks, Ourofino, farmacia veterinaria",
             "instrucao": "De 1 tablete a cada 24 horas por 90 dias."},
        ]}],
        [{"via": "USO TOPICO", "medicamentos": [
            {"item": "Pure+ Shampoo, Wesen Green, farmacia veterinaria",
             "instrucao": "Aplique nos pelos molhados..."},
        ]}],
    ]
    """
    doc = Document(docx_path)

    # Limpa o body preservando sectPr (margens + footer)
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    counter = 1

    for page_idx, vias in enumerate(pages):
        _presc_header(doc)
        _presc_line(doc)
        _presc_data(doc)
        _presc_line(doc)

        for via_dict in vias:
            via_name     = via_dict["via"]
            medicamentos = via_dict["medicamentos"]

            p_via = doc.add_paragraph()
            _pr(p_via, via_name, bold=True, size_pt=14, color=_PRESC_DARK_R)
            _para_space(p_via, before=200, after=100)

            for med in medicamentos:
                p_med = doc.add_paragraph()
                _pr(p_med, f"{counter}. {med['item']}",
                    bold=True, italic=True, size_pt=11, color=_PRESC_DARK_R)
                _para_space(p_med, before=80, after=60)

                p_instr = doc.add_paragraph()
                pPr = p_instr._element.get_or_add_pPr()
                ind = OxmlElement("w:ind")
                ind.set(qn("w:left"), "360")
                pPr.append(ind)
                _pr(p_instr, med["instrucao"], italic=True, size_pt=11,
                    color=_PRESC_DARK_R)
                _para_space(p_instr, before=0, after=120)

                counter += 1

        _add_signature_at_end(doc)

        if page_idx < len(pages) - 1:
            p_br = doc.add_paragraph()
            run_br = p_br.add_run()
            br = OxmlElement("w:br")
            br.set(qn("w:type"), "page")
            run_br._r.append(br)
            _para_space(p_br, before=0, after=0)

    _replace_all_placeholders(doc, data)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ═══════════════════════════════════════════════════════════════════════════════
# INSTRUÇÕES DIETA DE ELIMINAÇÃO
# ═══════════════════════════════════════════════════════════════════════════════

def build_eliminacao_instructions_docx(subtipo: str) -> bytes:
    """Gera o docx das instruções de dieta de eliminação.

    Parameters
    ----------
    subtipo : str
        'caseira' — 3 páginas (eliminação + preparo caseiro)
        'racao'   — 2 páginas (só eliminação)

    Returns
    -------
    bytes  — conteúdo do .docx pronto para convert_to_pdf()
    """
    from docx.shared import Cm

    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(3.5)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    TEAL = RGBColor(0x1A, 0x52, 0x76)
    DARK = RGBColor(0x1C, 0x28, 0x33)

    def _title(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(14)
        run.font.color.rgb = TEAL
        _para_space(p, before=0, after=200)

    def _subtitle(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = TEAL
        _para_space(p, before=200, after=120)

    def _intro(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
        _para_space(p, before=0, after=160)

    def _section(num, text):
        p = doc.add_paragraph()
        run = p.add_run(f"{num}. {text}")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = TEAL
        _para_space(p, before=160, after=60)

    def _bullet(text, level=1):
        p = doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(360 * level))
        ind.set(qn("w:hanging"), "180")
        pPr.append(ind)
        _para_space(p, before=0, after=60)
        char = "•" if level == 1 else "◦"
        run = p.add_run(f"{char} {text}")
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = DARK

    def _page_break():
        p = doc.add_paragraph()
        run = p.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._r.append(br)
        _para_space(p, before=0, after=0)

    # ── Página 1: seções 1–4 ──────────────────────────────────────────────────
    _title("INSTRUCOES SOBRE DIETA DE ELIMINACAO")
    _intro(
        "A dieta de eliminacao tem como objetivo identificar possiveis alimentos "
        "envolvidos em reacoes adversas, como alergias ou intolerancias alimentares."
    )

    _section(1, "Composicao da dieta")
    _bullet(
        "A dieta devera conter apenas uma fonte de proteina animal e uma fonte de "
        "carboidrato, que o animal nunca tenha consumido anteriormente OU uma dieta "
        "com proteina extensamente hidrolisada e inedita, conforme prescricao."
    )
    _bullet("Nao realizar substituicoes ou acrescimos sem orientacao profissional.")

    _section(2, "Exclusividade da dieta")
    _bullet("Durante o periodo da dieta de eliminacao, o animal NAO pode receber:")
    for item in [
        "Petiscos", "Ossos", "Biscoitos",
        "Frutas, legumes extras", "Restos de comida",
        "Suplementos, medicamentos palatáveis ou pastas com sabor",
    ]:
        _bullet(item, level=2)
    _bullet("Agua esta liberada a vontade.")

    _section(3, "Periodo minimo")
    _bullet("A dieta deve ser seguida de forma rigorosa por no minimo 8 semanas.")
    _bullet(
        "Melhoras parciais podem ocorrer antes, mas nao interromper a dieta "
        "antes do periodo recomendado."
    )

    _section(4, "Manejo alimentar")
    _bullet("Dividir a alimentacao diaria em 2 a 3 refeicoes, respeitando os horarios.")
    _bullet("Manter rotina alimentar estavel, sem variacoes de preparo ou ingredientes.")

    # ── Página 2: seções 5–6 + aviso ─────────────────────────────────────────
    _page_break()
    _title("INSTRUCOES SOBRE DIETA DE ELIMINACAO")

    _section(5, "Avaliacao clinica")
    _bullet(
        "Sinais como prurido, lesoes de pele, otites, vomitos ou diarreia devem "
        "ser monitorados e relatados."
    )
    _bullet(
        "Caso ocorra ingestao acidental de outro alimento, a dieta pode precisar "
        "ser reiniciada."
    )

    _section(6, "Fase de desafio alimentar")
    _bullet(
        "Apos o periodo de eliminacao e melhora clinica, novos alimentos so deverao "
        "ser reintroduzidos com orientacao profissional, um por vez, para identificacao "
        "do possivel agente causador."
    )

    p_warn = doc.add_paragraph()
    _para_space(p_warn, before=200, after=80)
    r_warn = p_warn.add_run(
        "O sucesso da dieta depende diretamente do cumprimento rigoroso das "
        "orientacoes. Qualquer excecao pode comprometer o diagnostico."
    )
    r_warn.bold = True
    r_warn.font.name = "Arial"
    r_warn.font.size = Pt(11)
    r_warn.font.color.rgb = DARK

    # ── Página 3 (apenas caseira): instruções gerais de preparo ───────────────
    if subtipo == "caseira":
        _page_break()
        _title("INSTRUCOES GERAIS DA DIETA")
        _subtitle("ORIENTACOES ESSENCIAIS PARA O PREPARO DA DIETA CASEIRA")
        _intro(
            "Esta dieta foi formulada especificamente para o seu pet. Para garantir "
            "que ele receba todos os nutrientes de forma segura e equilibrada, siga "
            "as orientacoes abaixo:"
        )

        _section(1, "Nao faca substituicoes ou alteracoes")
        _bullet("Nao troque ingredientes nem mude quantidades.")
        _bullet(
            "Toda a proporcao foi calculada individualmente. Ajustes sem orientacao "
            "podem desbalancear a dieta e prejudicar a saude do animal."
        )
        _bullet("Forneca o suplemento apenas no momento de servir o alimento.")

        _section(2, "Modo de preparo")
        _bullet(
            "Siga exatamente o indicado. A forma de preparo influencia diretamente "
            "na digestibilidade e aproveitamento dos nutrientes."
        )
        _bullet("Como pesar os ingredientes: pesar apos o preparo.")
        _bullet(
            "Cozido: Cozinhe em agua, sem oleo, sem sal e sem temperos. Nao utilize "
            "caldos prontos. Graos devem ficar 12 horas de molho antes de cozinhar."
        )
        _bullet("Assado: Forno ou airfryer ate 200 graus Celsius, sem oleo ou temperos.")
        _bullet(
            "Refogado: Use apenas uma pequena quantidade de agua para ajudar no "
            "cozimento. Nunca use oleo. Refogue por poucos minutos, com tampa, sem dourar."
        )

        _subtitle("HIGIENE E SEGURANCA ALIMENTAR")
        _bullet(
            "Evite contaminacao cruzada: lave bem maos, utensilios e superficies "
            "entre o manuseio de carnes cruas e outros alimentos."
        )
        _bullet("Armazenamento: Geladeira: ate 3 dias. Freezer: ate 90 dias.")
        _bullet(
            "Guarde sempre em recipientes limpos, bem vedados e porcoes individuais "
            "para facilitar o uso."
        )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
